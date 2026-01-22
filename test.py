# pip install python-docx
import copy
from decimal import Decimal, ROUND_HALF_UP
from config import Config
from tools import DatabaseTools as Tools
from datetime import datetime

from docx import Document


TEMPLATE_PATH = "templates/template.docx" 
OUTPUT_PATH = "КП_1v1.docx"             


def _replace_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    # ВАЖНО: заменяем по run'ам, чтобы не убить форматирование
    for run in paragraph.runs:
        for k, v in mapping.items():
            if k in run.text:
                run.text = run.text.replace(k, v)


def _replace_everywhere(doc: Document, mapping: dict[str, str]) -> None:
    # body paragraphs
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)

    # tables in body
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)

    # headers/footers
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            _replace_in_paragraph(p, mapping)
        for p in sec.footer.paragraphs:
            _replace_in_paragraph(p, mapping)


def _set_cell_text_preserve(cell, text: str) -> None:
    # Не используем cell.text = ..., иначе слетит формат.
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if p.runs:
        p.runs[0].text = text
        # очистим остальные run'ы, если есть
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def _remove_row(table, row_idx: int) -> None:
    table._tbl.remove(table.rows[row_idx]._tr)


def _fmt_money(v) -> str:
    # формат как в шаблоне: 1,57 (запятая) + знак ¥
    x = Decimal(str(v)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    s = f"{x:.2f}".replace(".", ",")
    return f"¥{s}"


def fill_doc_like_template(
    template_path: str,
    output_path: str,
) -> None:
    doc = Document(template_path)

    data = {
        # header
        "ООО «АЛЬФА КАППА ИНЖИНИРИНГ»": "ООО «АЛЬФА КАППА ИНЖИНИРИНГ»",
        "ИНН 9731121825; КПП 772901001": "ИНН 9731121825; КПП 772901001",
        "121471, Г.МОСКВА, УЛ., РЯБИНОВАЯ, Д.26 СТР.1, ПОМЕЩ.141": "121471, Г.МОСКВА, УЛ., РЯБИНОВАЯ, Д.26 СТР.1, ПОМЕЩ.141",
        "alphakappa.ru": "alphakappa.ru",
        "+7 (993) 338-47-22": "+7 (993) 338-47-22",
        "admin@alphakappa.ru": "admin@alphakappa.ru",

        # body/table 0
        "Исх. №9/19.01 от 19.01.2026": "Исх. №9/19.01 от 19.01.2026",
        "Директору": "Директору",
        "ООО Сусуман": "ООО Сусуман",
        "Иванов И. И.": "Земцов И. И.",

        # greeting + text
        "Уважаемая Иван Иваныч !": "Уважаемая Иван Иваныч !",
        "заявку 123": "заявку 123",

        # totals paragraph + terms
        "4,71 CNY": "4,71 CNY",
        "0,93 CNY": "0,93 CNY",
        "Срок гарантии - 123": "Срок гарантии - 123",
        "Производитель - 123;": "Производитель - 123;",
        "в течение 123": "в течение 123",
        "Срок действия КП до 29.01.2026.": "Срок действия КП до 29.01.2026.",

        # signature
        "Гениральнай директор": "Гениральнай директор",
        "А. О. Кадыров": "А. О. Кадыров",
    }

    # ====== 1) Глобальные замены (если ты хочешь менять поля) ======
    _replace_everywhere(doc, data)

    # ====== 2) Таблица с товарами: пересобрать строки по списку items, сохранив формат 1-в-1 ======
    # В твоем файле это doc.tables[1]
    products = doc.tables[1]

    # Сохраняем XML-шаблон одной товарной строки (строка 1)
    row_template_tr = copy.deepcopy(products.rows[1]._tr)

    # Удаляем все товарные строки между заголовком и итогом
    # Останутся: row 0 (шапка) и last row (итого)
    while len(products.rows) > 2:
        _remove_row(products, 1)

    # Товары (как в файле, чтобы вышло 1-в-1)
    items = [
        {"name": "Ось 2065214", "sku": "2065214", "unit": "шт.", "qty": 1, "price": "1.57", "days": "40 дней"},
        {"name": "Редуктор бортовой 1857593", "sku": "1857593", "unit": "шт.", "qty": 1, "price": "1.57", "days": "50 дней"},
        {"name": "Какашка 1857593", "sku": "1857593", "unit": "шт.", "qty": 1, "price": "1.57", "days": "70 дней"},
    ]

    totals_tr = products.rows[-1]._tr
    total_wo = Decimal("0.00")
    total_w = Decimal("0.00")

    for i, it in enumerate(items, start=1):
        new_tr = copy.deepcopy(row_template_tr)
        totals_tr.addprevious(new_tr)

        # только что вставленная строка всегда станет предпоследней
        r = products.rows[-2]

        price = Decimal(str(it["price"]))
        qty = Decimal(str(it["qty"]))
        row_wo = (price * qty).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        row_w = (row_wo * Decimal("1.20")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)

        total_wo += row_wo
        total_w += row_w

        _set_cell_text_preserve(r.cells[0], str(i))
        _set_cell_text_preserve(r.cells[1], it["name"])
        _set_cell_text_preserve(r.cells[2], it["sku"])
        _set_cell_text_preserve(r.cells[3], it["unit"])
        _set_cell_text_preserve(r.cells[4], str(it["qty"]))
        _set_cell_text_preserve(r.cells[5], _fmt_money(price))
        _set_cell_text_preserve(r.cells[6], _fmt_money(row_wo))
        _set_cell_text_preserve(r.cells[7], _fmt_money(row_w))
        _set_cell_text_preserve(r.cells[8], it["days"])

    # Итоговая строка (в шаблоне первые 6 колонок уже слиты)
    total_row = products.rows[-1]
    _set_cell_text_preserve(total_row.cells[0], "Итого")
    _set_cell_text_preserve(total_row.cells[6], _fmt_money(total_wo))
    _set_cell_text_preserve(total_row.cells[7], _fmt_money(total_w))

    doc.save(output_path)


if __name__ == "__main__":
    fill_doc_like_template('templates/template.docx', f"/КП_от_{datetime.now().strftime('%d.%m.%Y')}.docx")
    print(f"Saved: {OUTPUT_PATH}")
