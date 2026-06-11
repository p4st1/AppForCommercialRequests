import unittest

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from create import (
    FULL_PRODUCTS_TABLE_WIDTHS,
    MULTIPAGE_TABLE_TOP_MARGIN_PT,
    _apply_top_indent_for_multipage_table,
    _format_delivery_days_text,
    _format_delivery_period,
    _normalize_offer_terms_paragraphs,
    _normalize_products_table_headers,
    _optimize_products_table_layout,
    _section_text_width_dxa,
    _set_row_bold,
    _table_indent_dxa,
)


class CreateDocxLayoutTests(unittest.TestCase):
    def test_product_table_layout_fits_section_width_without_repeating_header(self):
        document = Document()
        table = document.add_table(rows=2, cols=len(FULL_PRODUCTS_TABLE_WIDTHS))
        for index, cell in enumerate(table.rows[0].cells):
            cell.text = f"h{index}"
        for index, cell in enumerate(table.rows[1].cells):
            cell.text = f"c{index}"

        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)

        _optimize_products_table_layout(
            table,
            document.sections[0],
            include_days=True,
            header_idx=0,
        )

        expected_width = _section_text_width_dxa(document.sections[0])
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        self.assertEqual(tbl_w.get(qn("w:type")), "dxa")
        self.assertEqual(int(tbl_w.get(qn("w:w"))), expected_width)

        grid_widths = [
            int(grid_col.get(qn("w:w")))
            for grid_col in table._tbl.tblGrid.gridCol_lst
        ]
        self.assertEqual(len(grid_widths), len(FULL_PRODUCTS_TABLE_WIDTHS))
        self.assertEqual(sum(grid_widths), expected_width)

        header = table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader"))
        self.assertIsNone(header)

    def test_product_table_layout_centers_table(self):
        document = Document()
        table = document.add_table(rows=1, cols=len(FULL_PRODUCTS_TABLE_WIDTHS))

        tbl_ind = OxmlElement("w:tblInd")
        tbl_ind.set(qn("w:w"), "432")
        tbl_ind.set(qn("w:type"), "dxa")
        table._tbl.tblPr.append(tbl_ind)

        _optimize_products_table_layout(
            table,
            document.sections[0],
            include_days=True,
            header_idx=0,
        )

        self.assertEqual(_table_indent_dxa(table), 0)
        table_alignment = table._tbl.tblPr.find(qn("w:jc"))
        self.assertIsNotNone(table_alignment)
        self.assertEqual(table_alignment.get(qn("w:val")), "center")

    def test_products_table_headers_use_requested_wording(self):
        document = Document()
        table = document.add_table(rows=1, cols=2)

        unit_paragraph = table.rows[0].cells[0].paragraphs[0]
        unit_paragraph.add_run("Ед. ")
        unit_paragraph.add_run("изм")
        price_paragraph = table.rows[0].cells[1].paragraphs[0]
        price_paragraph.add_run("Цена за ед. без ")
        price_paragraph.add_run("ндс")

        _normalize_products_table_headers(table, 0)

        self.assertEqual(table.rows[0].cells[0].text, "Ед. изм.")
        self.assertEqual(table.rows[0].cells[1].text, "Цена за ед. без НДС")

    def test_offer_terms_are_normalized_after_template_fill(self):
        document = Document()
        warranty = document.add_paragraph(
            "Срок гарантии - 12 месяцев, с момента отгрузки гарантия не "
            "распространяется на быстроизнашивающиеся части;"
        )
        payment = document.add_paragraph(
            "Оплата осуществляется в Рублях РФ по курсу Центрального банка РФ "
            "на дату оплаты."
        )

        _normalize_offer_terms_paragraphs(document)

        self.assertEqual(
            warranty.text,
            "Срок гарантии - 12 месяцев с момента поставки. "
            "Гарантия не распространяется на быстроизнашиваемые части;",
        )
        self.assertEqual(
            payment.text,
            "Оплата осуществляется в Рублях РФ по курсу Центрального банка РФ "
            "на дату оплаты;",
        )

    def test_set_row_bold_formats_every_total_cell_run(self):
        document = Document()
        table = document.add_table(rows=1, cols=3)
        table.rows[0].cells[0].text = "Итого"
        table.rows[0].cells[1].text = "100"
        table.rows[0].cells[2].text = "120"

        _set_row_bold(table.rows[0])

        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    self.assertTrue(run.bold)

    def test_apply_top_indent_for_multipage_table_uses_minimum_safe_margin(self):
        document = Document()
        document.sections[0].top_margin = Pt(56.7)
        items = [
            {
                "name": "Длинное наименование позиции " * 3,
                "sku": "1234567890",
                "days": "130 дней",
            }
            for _ in range(18)
        ]

        _apply_top_indent_for_multipage_table(
            document,
            items,
            include_days=True,
        )

        self.assertGreaterEqual(
            round(document.sections[0].top_margin.pt, 1),
            float(MULTIPAGE_TABLE_TOP_MARGIN_PT),
        )

    def test_delivery_days_text_blanks_zero_days(self):
        self.assertEqual(_format_delivery_days_text("0 дней"), "")
        self.assertEqual(_format_delivery_days_text(" 80 дней "), "80 дней")

    def test_delivery_period_ignores_zero_days(self):
        self.assertEqual(
            _format_delivery_period(["0 дней", "80 дней", "130 дней"]),
            "от 80 до 130",
        )
        self.assertEqual(_format_delivery_period(["0 дней", "130 дней"]), "до 130")
        self.assertEqual(_format_delivery_period(["0 дней", ""]), "")


if __name__ == "__main__":
    unittest.main()
