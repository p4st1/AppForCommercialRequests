from __future__ import annotations

import csv
import json
import re
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Any


SUBMISSION_HEADERS = (
    "Наименование",
    "Кол-во",
    "Ед. изм.",
    "Цена за ед.",
    "Сумма",
    "Срок поставки",
    "Производитель",
    "Технические характеристики",
    "Статус поставщика",
    "Гарантия",
)


FIELD_ORDER = (
    "name",
    "qty",
    "unit",
    "unit_price",
    "total",
    "delivery_time",
    "manufacturer",
    "technical",
    "supplier_status",
    "warranty",
)


FIELD_LABELS = dict(zip(FIELD_ORDER, SUBMISSION_HEADERS, strict=True))


@dataclass
class SubmissionHeader:
    number: str = ""
    title: str = ""
    customer: str = ""
    currency: str = ""
    offer_validity_period: str = ""
    delivery_order: str = ""
    payment_terms: str = ""
    payment_condition: str = ""
    total: float = 0.0
    lot_id: str = ""


@dataclass
class SubmissionRow:
    name: str = ""
    qty: float | None = None
    unit: str = ""
    unit_price: float | None = None
    total: float | None = None
    delivery_time: str = ""
    manufacturer: str = ""
    technical: str = ""
    supplier_status: str = ""
    warranty: str = ""

    def has_content(self) -> bool:
        return any(
            str(value or "").strip()
            for value in (
                self.name,
                self.qty,
                self.unit,
                self.unit_price,
                self.total,
                self.delivery_time,
                self.manufacturer,
                self.technical,
                self.supplier_status,
                self.warranty,
            )
        )

    def to_cells(self) -> list[Any]:
        return [getattr(self, field) for field in FIELD_ORDER]


@dataclass
class SubmissionPayload:
    header: SubmissionHeader
    rows: list[SubmissionRow]

    def to_dict(self) -> dict[str, Any]:
        return {
            "header": asdict(self.header),
            "rows": [asdict(row) for row in self.rows],
        }


@dataclass
class SubmissionValidationIssue:
    row: int | None
    field: str
    label: str
    severity: str
    message: str


class SubmissionService:
    CURRENCY_CODES = {
        "AUD",
        "BYN",
        "CHF",
        "CNY",
        "EUR",
        "GBP",
        "INR",
        "JPY",
        "KZT",
        "RSD",
        "RUB",
        "TRY",
        "UAH",
        "USD",
    }
    CURRENCY_ALIASES = {
        "₽": "RUB",
        "руб": "RUB",
        "рубл": "RUB",
        "rur": "RUB",
        "$": "USD",
        "доллар": "USD",
        "usd": "USD",
        "€": "EUR",
        "евро": "EUR",
        "eur": "EUR",
        "¥": "CNY",
        "юан": "CNY",
        "yuan": "CNY",
        "cny": "CNY",
        "cyn": "CNY",
        "тенге": "KZT",
        "kzt": "KZT",
    }
    REQUIRED_HEADER_FIELDS = {
        "number": "Номер заявки",
        "title": "Название заявки",
        "offer_validity_period": "Срок действия КП",
    }
    REQUIRED_ROW_FIELDS = {
        "name": "Наименование",
        "qty": "Кол-во",
        "unit_price": "Цена за ед.",
        "delivery_time": "Срок поставки",
        "manufacturer": "Производитель",
        "technical": "Технические характеристики",
    }
    WARNING_ROW_FIELDS = {
        "unit": "Ед. изм.",
        "supplier_status": "Статус поставщика",
        "warranty": "Гарантия",
    }

    @classmethod
    def normalize_currency_code(cls, value: Any) -> str:
        text = str(value or "").strip()
        if not text:
            return ""

        upper_text = text.upper()
        for code in sorted(cls.CURRENCY_CODES, key=len, reverse=True):
            if re.search(rf"(?<![A-Z]){re.escape(code)}(?![A-Z])", upper_text):
                return code

        normalized = text.casefold().replace("ё", "е")
        compact = re.sub(r"\s+", " ", normalized)
        for alias, code in cls.CURRENCY_ALIASES.items():
            if alias in compact:
                return code
        return ""

    @classmethod
    def _currency_from_matrix(cls, matrix: list[list[Any]]) -> str:
        currency_columns = []
        plain_cells = []
        for row in matrix[:30]:
            for cell in list(row or [])[:80]:
                text = str(cell or "").strip()
                if not text:
                    continue
                normalized = cls._normalize_header(text)
                if (
                    "валют" in normalized
                    or "сумма" in normalized
                    or "цена" in normalized
                    or "стоим" in normalized
                ):
                    currency_columns.append(text)
                else:
                    plain_cells.append(text)

        for text in currency_columns + plain_cells:
            code = cls.normalize_currency_code(text)
            if code:
                return code
        return ""

    @classmethod
    def detect_currency(cls, file_path: str | Path) -> str:
        path = Path(file_path).expanduser()
        if not path.exists() or not path.is_file():
            return ""

        suffix = path.suffix.lower()
        try:
            if suffix == ".docx":
                from docx import Document

                document = Document(str(path))
                matrix: list[list[Any]] = []
                for table in document.tables:
                    matrix.extend([[cell.text for cell in row.cells] for row in table.rows])
                return cls._currency_from_matrix(matrix)

            if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
                from openpyxl import load_workbook

                workbook = load_workbook(path, read_only=True, data_only=True)
                try:
                    worksheet = workbook.active
                    max_row = min(int(worksheet.max_row or 1), 30)
                    max_column = min(int(worksheet.max_column or 1), 80)
                    matrix = []
                    for row_index in range(1, max_row + 1):
                        row_values = []
                        for column_index in range(1, max_column + 1):
                            cell = worksheet.cell(
                                row=row_index,
                                column=column_index,
                            )
                            row_values.append(cell.value)
                            number_format = str(cell.number_format or "")
                            if number_format and number_format != "General":
                                row_values.append(number_format)
                        matrix.append(row_values)
                    return cls._currency_from_matrix(matrix)
                finally:
                    workbook.close()

            if suffix == ".csv":
                try:
                    text = path.read_text(encoding="utf-8-sig")
                except UnicodeDecodeError:
                    text = path.read_text(encoding="cp1251")
                sample = text[:4096]
                try:
                    dialect = csv.Sniffer().sniff(sample, delimiters=";,\t")
                    delimiter = dialect.delimiter
                except csv.Error:
                    delimiter = ";" if sample.count(";") >= sample.count(",") else ","
                matrix = list(csv.reader(text.splitlines()[:30], delimiter=delimiter))
                return cls._currency_from_matrix(matrix)

            if suffix == ".xls":
                import pandas as pd

                dataframe = pd.read_excel(path, header=None, nrows=30)
                return cls._currency_from_matrix(dataframe.values.tolist())
        except Exception:
            return ""
        return ""

    @staticmethod
    def parse_number(value: Any) -> float | None:
        if value is None or isinstance(value, bool):
            return None
        if isinstance(value, (int, float)):
            return float(value)

        text = (
            str(value)
            .strip()
            .replace("\xa0", " ")
            .replace(" ", "")
            .replace("₽", "")
            .replace("руб.", "")
            .replace("руб", "")
            .replace("RUB", "")
            .replace("rub", "")
            .replace("$", "")
            .replace("€", "")
            .replace("¥", "")
            .replace(",", ".")
        )
        if not text:
            return None
        if re.fullmatch(r"[-+]?\d+(?:\.\d+)?", text) is None:
            return None
        try:
            return float(text)
        except ValueError:
            return None

    @classmethod
    def format_money(cls, value: Any) -> str:
        number = cls.parse_number(value)
        if number is None:
            return "" if value is None else str(value)
        return f"{number:,.2f}".replace(",", " ").replace(".", ",")

    @classmethod
    def calculate_row_total(cls, row: SubmissionRow) -> float | None:
        explicit_total = cls.parse_number(row.total)
        if explicit_total is not None:
            return round(explicit_total, 2)

        qty = cls.parse_number(row.qty)
        price = cls.parse_number(row.unit_price)
        if qty is None or price is None:
            return None
        return round(qty * price, 2)

    @classmethod
    def normalize_row(cls, row: SubmissionRow) -> SubmissionRow:
        qty = cls.parse_number(row.qty)
        unit_price = cls.parse_number(row.unit_price)
        normalized = SubmissionRow(
            name=str(row.name or "").strip(),
            qty=qty,
            unit=str(row.unit or "").strip(),
            unit_price=unit_price,
            total=row.total,
            delivery_time=str(row.delivery_time or "").strip(),
            manufacturer=str(row.manufacturer or "").strip(),
            technical=str(row.technical or "").strip(),
            supplier_status=str(row.supplier_status or "").strip(),
            warranty=str(row.warranty or "").strip(),
        )
        normalized.total = cls.calculate_row_total(normalized)
        return normalized

    @classmethod
    def build_row_from_cells(cls, cells: list[Any] | tuple[Any, ...]) -> SubmissionRow:
        values = list(cells)[: len(FIELD_ORDER)]
        values.extend("" for _ in range(len(FIELD_ORDER) - len(values)))
        data = dict(zip(FIELD_ORDER, values, strict=True))
        return cls.normalize_row(SubmissionRow(**data))

    @classmethod
    def prepare_payload(
        cls,
        header: SubmissionHeader,
        rows: list[SubmissionRow],
    ) -> SubmissionPayload:
        normalized_rows = [
            cls.normalize_row(row)
            for row in rows
            if isinstance(row, SubmissionRow) and row.has_content()
        ]
        total = sum(
            float(row.total or 0.0)
            for row in normalized_rows
            if cls.parse_number(row.total) is not None
        )
        normalized_header = SubmissionHeader(
            number=str(header.number or "").strip(),
            title=str(header.title or "").strip(),
            customer=str(header.customer or "").strip(),
            currency=cls.normalize_currency_code(header.currency)
            or str(header.currency or "").strip(),
            offer_validity_period=str(
                getattr(header, "offer_validity_period", "") or ""
            ).strip(),
            delivery_order=str(getattr(header, "delivery_order", "") or "").strip(),
            payment_terms=str(getattr(header, "payment_terms", "") or "").strip(),
            payment_condition=str(
                getattr(header, "payment_condition", "") or ""
            ).strip(),
            lot_id=str(getattr(header, "lot_id", "") or "").strip(),
            total=round(total, 2),
        )
        return SubmissionPayload(header=normalized_header, rows=normalized_rows)

    @classmethod
    def validate(
        cls,
        header: SubmissionHeader,
        rows: list[SubmissionRow],
    ) -> list[SubmissionValidationIssue]:
        payload = cls.prepare_payload(header, rows)
        issues: list[SubmissionValidationIssue] = []

        for field, label in cls.REQUIRED_HEADER_FIELDS.items():
            if not str(getattr(payload.header, field, "") or "").strip():
                issues.append(
                    SubmissionValidationIssue(
                        row=None,
                        field=field,
                        label=label,
                        severity="error",
                        message=f"Заполните поле '{label}'",
                    )
                )

        if not payload.rows:
            issues.append(
                SubmissionValidationIssue(
                    row=None,
                    field="rows",
                    label="Позиции",
                    severity="error",
                    message="Добавьте хотя бы одну позицию",
                )
            )
            return issues

        for row_index, row in enumerate(payload.rows):
            for field, label in cls.REQUIRED_ROW_FIELDS.items():
                value = getattr(row, field)
                if field in {"qty", "unit_price"}:
                    if cls.parse_number(value) is None:
                        issues.append(
                            SubmissionValidationIssue(
                                row=row_index,
                                field=field,
                                label=label,
                                severity="error",
                                message=f"Строка {row_index + 1}: заполните '{label}'",
                            )
                        )
                    continue
                if not str(value or "").strip():
                    issues.append(
                        SubmissionValidationIssue(
                            row=row_index,
                            field=field,
                            label=label,
                            severity="error",
                            message=f"Строка {row_index + 1}: заполните '{label}'",
                        )
                    )

            for field, label in cls.WARNING_ROW_FIELDS.items():
                if not str(getattr(row, field) or "").strip():
                    issues.append(
                        SubmissionValidationIssue(
                            row=row_index,
                            field=field,
                            label=label,
                            severity="warning",
                            message=f"Строка {row_index + 1}: проверьте '{label}'",
                        )
                    )
        return issues

    @staticmethod
    def has_errors(issues: list[SubmissionValidationIssue]) -> bool:
        return any(issue.severity == "error" for issue in issues)

    @staticmethod
    def _normalize_header(value: Any) -> str:
        text = str(value or "").strip().casefold().replace("ё", "е")
        return re.sub(r"[^a-zа-я0-9]+", "", text)

    @classmethod
    def _column_kind(cls, header: Any) -> str | None:
        candidates = cls._column_candidates(header)
        if not candidates:
            return None
        return max(candidates, key=lambda item: item[1])[0]

    @classmethod
    def _column_candidates(cls, header: Any) -> list[tuple[str, int]]:
        normalized = cls._normalize_header(header)
        if not normalized:
            return []

        candidates: list[tuple[str, int]] = []
        if normalized in {"наименование", "наименованиетовара", "name", "title"}:
            candidates.append(("name", 100))
        elif "альтернативноенаимен" in normalized:
            candidates.append(("name", 20))
        elif "наимен" in normalized:
            candidates.append(("name", 40))
        if (
            "предлагаемоеколво" in normalized
            or "предлагаемоеколичество" in normalized
        ):
            candidates.append(("qty", 100))
        elif "колво" in normalized or "количество" in normalized or normalized == "qty":
            candidates.append(("qty", 80))
        if "едизм" in normalized or "единиц" in normalized or normalized == "unit":
            candidates.append(("unit", 100))
        if (
            "предлагаемаяценазаед" in normalized
            and "сумма" not in normalized
            and "итого" not in normalized
        ):
            candidates.append(("unit_price", 100))
        elif "цена" in normalized and "сумма" not in normalized and "итого" not in normalized:
            candidates.append(("unit_price", 70))
        if "сумма" in normalized or "итого" in normalized or normalized == "total":
            candidates.append(("total", 90))
        if "срок" in normalized and "постав" in normalized:
            candidates.append(("delivery_time", 100))
        if "производ" in normalized:
            candidates.append(("manufacturer", 100))
        if "характерист" in normalized or "техническ" in normalized:
            candidates.append(("technical", 100))
        if "статус" in normalized and "постав" in normalized:
            candidates.append(("supplier_status", 100))
        if "гарант" in normalized:
            candidates.append(("warranty", 100))
        return candidates

    @classmethod
    def _rows_from_matrix(cls, matrix: list[list[Any]]) -> list[SubmissionRow]:
        cleaned = [
            cls._split_single_cell_delimited_row(list(row or []))
            for row in matrix
            if any(str(cell or "").strip() for cell in row or [])
        ]
        if not cleaned:
            return []

        header_index = 0
        column_map: dict[int, str] = {}
        for row_index, row in enumerate(cleaned[:10]):
            detected_fields: dict[str, tuple[int, int]] = {}
            for column_index, cell in enumerate(row):
                for kind, priority in cls._column_candidates(cell):
                    current = detected_fields.get(kind)
                    if current is None or priority > current[1]:
                        detected_fields[kind] = (column_index, priority)
            detected = {
                column_index: kind
                for kind, (column_index, _priority) in detected_fields.items()
            }
            if len(detected) >= 2:
                header_index = row_index
                column_map = detected
                break

        if not column_map:
            column_map = dict(enumerate(FIELD_ORDER))
            data_rows = cleaned
        else:
            data_rows = cleaned[header_index + 1 :]

        rows: list[SubmissionRow] = []
        for source_row in data_rows:
            values = {field: "" for field in FIELD_ORDER}
            for column_index, field in column_map.items():
                if column_index < len(source_row):
                    values[field] = source_row[column_index]
            row = cls.normalize_row(SubmissionRow(**values))
            if row.has_content():
                rows.append(row)
        return rows

    @staticmethod
    def _split_single_cell_delimited_row(row: list[Any]) -> list[Any]:
        if len(row) != 1:
            return row
        cell = str(row[0] or "").strip()
        if not cell:
            return row
        delimiter = ";" if ";" in cell else "\t" if "\t" in cell else ""
        if not delimiter:
            return row
        try:
            parsed = next(csv.reader([cell], delimiter=delimiter))
        except Exception:
            return row
        return parsed if len(parsed) > 1 else row

    @classmethod
    def load_kp(cls, file_path: str | Path) -> list[SubmissionRow]:
        path = Path(file_path).expanduser()
        if not path.exists() or not path.is_file():
            raise FileNotFoundError(f"Файл КП не найден: {path}")

        suffix = path.suffix.lower()
        if suffix == ".docx":
            return cls._load_docx(path)
        if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
            return cls._load_xlsx(path)
        if suffix == ".xls":
            return cls._load_xls(path)
        if suffix == ".csv":
            return cls._load_csv(path)
        raise ValueError("Поддерживаются КП в форматах .docx, .xlsx, .xlsm, .xls, .csv")

    @classmethod
    def _load_docx(cls, path: Path) -> list[SubmissionRow]:
        from docx import Document

        document = Document(str(path))
        all_rows: list[SubmissionRow] = []
        for table in document.tables:
            matrix = [[cell.text for cell in row.cells] for row in table.rows]
            all_rows.extend(cls._rows_from_matrix(matrix))
        return all_rows

    @classmethod
    def _load_xlsx(cls, path: Path) -> list[SubmissionRow]:
        from openpyxl import load_workbook

        workbook = load_workbook(str(path), data_only=True)
        try:
            worksheet = workbook.active
            matrix = [list(row) for row in worksheet.iter_rows(values_only=True)]
            return cls._rows_from_matrix(matrix)
        finally:
            workbook.close()

    @classmethod
    def _load_xls(cls, path: Path) -> list[SubmissionRow]:
        import pandas as pd

        dataframe = pd.read_excel(str(path), header=None, dtype=object).fillna("")
        matrix = dataframe.values.tolist()
        return cls._rows_from_matrix(matrix)

    @classmethod
    def _load_csv(cls, path: Path) -> list[SubmissionRow]:
        with path.open("r", encoding="utf-8-sig", newline="") as file_obj:
            sample = file_obj.read(4096)
            file_obj.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=",;\t")
            except csv.Error:
                dialect = csv.excel
                if sample.count(";") > sample.count(","):
                    dialect.delimiter = ";"
            rows = list(csv.reader(file_obj, dialect))
        return cls._rows_from_matrix(rows)

    @staticmethod
    def _safe_file_part(value: Any) -> str:
        text = str(value or "").strip()
        text = re.sub(r"[^0-9A-Za-zА-Яа-я_-]+", "_", text).strip("_")
        return text or "submission"

    @classmethod
    def save_payload(
        cls,
        payload: SubmissionPayload,
        directory: str | Path,
    ) -> Path:
        target_dir = Path(directory).expanduser()
        target_dir.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        number = cls._safe_file_part(payload.header.number)
        target_path = target_dir / f"submission_{number}_{timestamp}.json"
        target_path.write_text(
            json.dumps(payload.to_dict(), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        return target_path.resolve()
