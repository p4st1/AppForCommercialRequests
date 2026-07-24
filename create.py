from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from decimal import Decimal, ROUND_HALF_UP
from tools import DatabaseTools as Tools
from tools import Tools as ExtraTools
from datetime import datetime, timedelta
import copy
import math
from openpyxl import load_workbook
from openpyxl.styles import Border, Side, Alignment
from config import Config
from services.excel_recalc import force_excel_recalc
import shutil
import subprocess
import os
import re
from pathlib import Path


FULL_PRODUCTS_TABLE_WIDTHS = (420, 1800, 1550, 560, 700, 1080, 1220, 1220, 1139)
SHORT_PRODUCTS_TABLE_WIDTHS = (430, 2600, 1600, 600, 650, 1250, 1279, 1279)
MULTIPAGE_TABLE_TOP_MARGIN_PT = 92
PRODUCT_CONDITION_TEXT = (
    "Продукция новая, не бывшая ранее в использовании. "
    "Год производства не позднее 2025;"
)
WARRANTY_CLAUSE_SUFFIX = (
    " с момента поставки. Гарантия не распространяется на быстроизнашиваемые части;"
)


def _length_to_dxa(length) -> int:
    return int(round(int(length) / 635))


def _section_text_width_dxa(section) -> int:
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    return max(1, _length_to_dxa(page_width - left_margin - right_margin))


def _scaled_widths(base_widths: tuple[int, ...], target_width: int) -> list[int]:
    if target_width <= 0:
        return list(base_widths)

    base_total = sum(base_widths)
    if base_total <= 0:
        return list(base_widths)

    scaled = [max(1, int(round(width * target_width / base_total))) for width in base_widths]
    delta = target_width - sum(scaled)
    if delta:
        widest_idx = max(range(len(scaled)), key=scaled.__getitem__)
        scaled[widest_idx] = max(1, scaled[widest_idx] + delta)
    return scaled


def _set_or_append(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _set_paragraph_text_keep_style(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return

    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def _set_cell_text_keep_style(cell, text: str) -> None:
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    _set_paragraph_text_keep_style(p, text)
    for extra_p in cell.paragraphs[1:]:
        for r in extra_p.runs:
            r.text = ""


def _canonical_docx_text(value) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip().casefold()


def _normalize_products_table_headers(table, header_idx: int) -> None:
    if not (0 <= header_idx < len(table.rows)):
        return

    replacements = {
        "ед. изм": "Ед. изм.",
        "ед. изм.": "Ед. изм.",
        "цена за ед. без ндс": "Цена за ед. без НДС",
    }
    for cell in table.rows[header_idx].cells:
        replacement = replacements.get(_canonical_docx_text(cell.text))
        if replacement:
            _set_cell_text_keep_style(cell, replacement)


def _table_indent_dxa(table) -> int:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return 0
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        return 0
    try:
        return int(tbl_ind.get(qn("w:w"), "0"))
    except (TypeError, ValueError):
        return 0


def _center_table(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    jc = _set_or_append(tbl_pr, "w:jc")
    jc.set(qn("w:val"), "center")

    tbl_ind = _set_or_append(tbl_pr, "w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "0")


def _normalize_warranty_clause_text(text: str) -> str:
    value = str(text or "")
    if "срок гарантии" not in value.casefold():
        return value

    match = re.match(r"^(.*?срок\s+гарантии\s*[-:]\s*)(.*)$", value, flags=re.IGNORECASE)
    if not match:
        return value

    prefix, rest = match.groups()
    period = re.split(
        r",?\s*с\s+момента\s+(?:отгрузки|поставки)\b",
        rest,
        maxsplit=1,
        flags=re.IGNORECASE,
    )[0]
    period = period.strip().rstrip(" ,.;")
    if not period:
        return f"{prefix.rstrip()}{WARRANTY_CLAUSE_SUFFIX}"
    return f"{prefix}{period}{WARRANTY_CLAUSE_SUFFIX}"


def _normalize_payment_clause_text(text: str) -> str:
    value = str(text or "")
    if "оплата осуществляется" not in value.casefold():
        return value
    return value.rstrip().rstrip(".,;") + ";"


def _normalize_offer_terms_paragraphs(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        original = paragraph.text
        normalized = _normalize_warranty_clause_text(original)
        normalized = _normalize_payment_clause_text(normalized)
        if normalized != original:
            _set_paragraph_text_keep_style(paragraph, normalized)


def _cell_grid_span(tc) -> int:
    tc_pr = tc.tcPr
    if tc_pr is None:
        return 1
    grid_span = tc_pr.find(qn("w:gridSpan"))
    if grid_span is None:
        return 1
    try:
        return max(1, int(grid_span.get(qn("w:val"), "1")))
    except (TypeError, ValueError):
        return 1


def _set_tc_width(tc, width: int) -> None:
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(max(1, int(width))))


def _set_table_grid_widths(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    total_width = sum(widths)
    tbl_w = _set_or_append(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_width))

    tbl_layout = _set_or_append(tbl_pr, "w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")

    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    insert_idx = 1 if tbl.tblPr is not None else 0
    tbl.insert(insert_idx, tbl_grid)

    for row in table.rows:
        col_idx = 0
        for tc in row._tr.tc_lst:
            span = _cell_grid_span(tc)
            width = sum(widths[col_idx: col_idx + span]) if col_idx < len(widths) else widths[-1]
            _set_tc_width(tc, width)
            col_idx += span


def _clear_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is not None:
        tr_pr.remove(tbl_header)


def _set_row_bold(row) -> None:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            if not paragraph.runs:
                paragraph.add_run("")
            for run in paragraph.runs:
                run.bold = True


def _insert_product_condition_after_delivery_terms(
    doc: Document,
    text: str = PRODUCT_CONDITION_TEXT,
) -> Paragraph | None:
    if any(_canonical_docx_text(paragraph.text) == _canonical_docx_text(text) for paragraph in doc.paragraphs):
        return None

    for paragraph in doc.paragraphs:
        if "условия поставки" not in _canonical_docx_text(paragraph.text):
            continue
        new_p = copy.deepcopy(paragraph._p)
        paragraph._p.addnext(new_p)
        new_paragraph = Paragraph(new_p, paragraph._parent)
        _set_paragraph_text_keep_style(new_paragraph, text)
        return new_paragraph
    return None


def _optimize_products_table_layout(table, section, *, include_days: bool, header_idx: int = 0) -> None:
    base_widths = FULL_PRODUCTS_TABLE_WIDTHS if include_days else SHORT_PRODUCTS_TABLE_WIDTHS
    widths = _scaled_widths(base_widths, _section_text_width_dxa(section))
    table.autofit = False
    _set_table_grid_widths(table, widths)
    _center_table(table)
    if 0 <= header_idx < len(table.rows):
        _clear_repeat_table_header(table.rows[header_idx])


def _parse_delivery_days(value) -> int:
    match = re.search(r"-?\d+", str(value or ""))
    if match is None:
        return 0
    try:
        return int(match.group())
    except ValueError:
        return 0


def _format_delivery_days_text(value) -> str:
    return str(value or "").strip() if _parse_delivery_days(value) > 0 else ""


def _format_delivery_period(values) -> str:
    days_values = [_parse_delivery_days(value) for value in values]
    positive_days = [days for days in days_values if days > 0]
    if not positive_days:
        return ""

    min_days = min(positive_days)
    max_days = max(positive_days)
    if min_days == max_days:
        return f"до {min_days}"
    return f"от {min_days} до {max_days}"


def _estimate_table_visual_rows(items: list[dict], *, include_days: bool) -> int:
    if not items:
        return 0

    name_chars_per_line = 34 if include_days else 42
    sku_chars_per_line = 24 if include_days else 28
    days_chars_per_line = 11

    visual_rows = 0
    for item in items:
        name_lines = max(1, math.ceil(len(str(item.get("name", "")).strip()) / name_chars_per_line))
        sku_lines = max(1, math.ceil(len(str(item.get("sku", "")).strip()) / sku_chars_per_line))
        days_lines = 1
        if include_days:
            days_lines = max(1, math.ceil(len(str(item.get("days", "")).strip()) / days_chars_per_line))
        visual_rows += max(name_lines, sku_lines, days_lines)
    return visual_rows


def _table_spans_multiple_pages(items: list[dict], *, include_days: bool) -> bool:
    capacity = 15 if include_days else 20
    return _estimate_table_visual_rows(items, include_days=include_days) > capacity


def _apply_top_indent_for_multipage_table(doc: Document, items: list[dict], *, include_days: bool) -> None:
    if not _table_spans_multiple_pages(items, include_days=include_days):
        return

    for section in doc.sections:
        current_margin = section.top_margin if section.top_margin is not None else Pt(0)
        if current_margin < Pt(MULTIPAGE_TABLE_TOP_MARGIN_PT):
            section.top_margin = Pt(MULTIPAGE_TABLE_TOP_MARGIN_PT)


class createTextFile:
    def __init__(self, docxData):
        self.output_path = ""
        self.success = False
        self.error_message = ""

        Tools.write_log(f"test feature: {Config.settings['testFeature']}")
        documents_dir = Path.home() / "Documents"
        docx_output_dir = Tools.ensure_directory(
            Config.config.get("pathToSaveCP"),
            documents_dir,
        )
        Config.config["pathToSaveCP"] = str(docx_output_dir)
        Tools.write_log(f"Docx path to save: {docx_output_dir}")
        Tools.write_log('INIT DOCX...')

        tableData = docxData[0][1]
        customerData = docxData[1][0]
        extraData = docxData[2]
        dbData = docxData[3]
        lastCol = docxData[4]
        condData = docxData[5]
        output_format = self._normalize_output_format(
            docxData[6] if len(docxData) > 6 else "docx"
        )

        sum1, sum2 = Decimal("0.00"), Decimal("0.00")

        tool = ExtraTools()

        delivery_days_values = []

        for i in tableData:
            currency1, amount1 = Tools.parsePrice(i[6].replace(',', '.'))
            currency2, amount2 = Tools.parsePrice(i[7].replace(',', '.'))
            sum1 += Decimal(amount1.replace(' ', '').replace(',', '.'))
            sum2 += Decimal(amount2.replace(' ', '').replace(',', '.'))
            symbCurrency = currency1
            delivery_days_values.append(i[8])

        def _round_money(v: Decimal) -> Decimal:
            return Decimal(str(v)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)

        period = _format_delivery_period(delivery_days_values)

        currency = Config.currency[symbCurrency]

        if docxData[4]:
            TEMPLATE_PATH = Config.template_docx_path
        else:
            TEMPLATE_PATH = Config.template_docx_path_short

        lot_number = str(extraData[0]).strip() if len(extraData) > 0 else ""
        today_date = datetime.now().strftime('%d_%m_%Y')
        OUTPUT_PATH = str(docx_output_dir / f"КП_{lot_number}_{today_date}.docx")
        self.output_path = OUTPUT_PATH

        VAT_RATE = Decimal("0.20")

        GENDER = {
            'мужской': ("ый", 'у'),
            'женский': ("ая", 'ой')
        }

        if Config.settings['testFeature']:
            post = Tools.formWord(customerData[8], 2)
            initials = f"{Tools.formWord(customerData[2], 2)} {customerData[1][0]}. {customerData[3][0]}."
        else:
            post = customerData[8]
            initials = f"{customerData[2]} {customerData[1][0]}. {customerData[3][0]}."

        PLACEHOLDERS = {
            "<<COMPANY>>": "ООО «АЛЬФА КАППА ИНЖИНИРИНГ»",
            "<<INN>>": "9731121825",
            "<<KPP>>": "772901001",
            "<<ADDRESS>>": "121471, Г.МОСКВА, УЛ., РЯБИНОВАЯ, Д.26 СТР.1, ПОМЕЩ.141",
            "<<SITE>>": "alphakappa.ru",
            "<<PHONE>>": "+7 (993) 338-47-22",
            "<<EMAIL>>": "admin@alphakappa.ru",

            "<<OUT_NUM>>": "9/19.01",
            "<<OUT_DATE>>": "19.01.2026",

            "<<TO_TITLE>>": "Директору",
            "<<CLIENT_COMPANY>>": "ООО Сусуман",
            "<<CLIENT_PERSON>>": "Иванов И. И.",

            "<<VALID_UNTIL>>": "29.01.2026",

            "<<SIGN_TITLE>>": "Гениральнай директор",
            "<<SIGN_NAME>>": "А. О. Кадыров",

            "{{num}}": f"{docxData[3]}",
            "{{date_f}}": f"{datetime.now().strftime('%d.%m')}",
            "{{now}}": f"{datetime.now().strftime('%d.%m.%Y')}",
            "{{post}}": post,
            "{{company}}": customerData[7],
            "{{initials}}": initials,

            "Уважаемая Иван Иваныч !": f"Уважаем{GENDER[customerData[10]][0]} {customerData[1]} {customerData[3]} !",
            "{{gender}}": f"{GENDER[customerData[10]][0]}",
            "{{name}}": f"{customerData[1]}",
            "{{surname}}": f"{customerData[3]}",
            "{{app_num}}": f"{extraData[0]}",

            "{{Total_wo}}": f"{Tools.num2text(_round_money(sum2))}",
            "{{Currency}}": f"{currency[0]}",
            "{{Total_diff}}": f"{Tools.num2text(_round_money(sum2 - sum1))}",

            "{{Total_wo_text}}": f"{tool.decimal2text(_round_money(sum2),int_units=currency[1],exp_units=currency[2])}",

            "{{Total_diff_text}}": f"{tool.decimal2text(_round_money(sum2 - sum1),int_units=currency[1],exp_units=currency[2])}",

            "{{NDS}}": f"{Tools.load_json(Config.vars_path)['parameters']['1'][1]}",

            "{{Conditions}}": f"{customerData[9]}",
            "{{Garanty_period}}": f"{extraData[1]}",
            "{{MinDays}}": f"{period}",
            "{{Producer}}": f"{extraData[3]}",
            "{{Pay_period}}": f"{extraData[2]}",
            "{{Pay_cond}}": f"{docxData[5]}",
            "{{Delivery_order}}": f"{extraData[5] if len(extraData) > 5 else ''}",
            "<<DELIVERY_ORDER>>": f"{extraData[5] if len(extraData) > 5 else ''}",
            "{{date_20days}}": f"{(datetime.now() + timedelta(days=Config.get_offer_validity_days())).strftime('%d.%m.%Y')}"
        }

        ITEMS = [

        ]

        for item in tableData:
            ITEMS.append({"name": item[1],
                          "sku": item[2],
                          "unit": item[3],
                          "qty": item[4],
                          "price": Decimal(Tools.parsePrice(item[5])[1].replace(',', '.').replace(' ', '')),
                          "sum_wo": Decimal(Tools.parsePrice(item[6])[1].replace(',', '.').replace(' ', '')),
                          "sum_w": Decimal(Tools.parsePrice(item[7])[1].replace(',', '.').replace(' ', '')),
                          "days": _format_delivery_days_text(item[8])})

        ROW_TOKENS = {
            "<<I>>": None,
            "<<Name>>": None,
            "<<SKU>>": None,
            "<<Unit>>": None,
            "<<Qty>>": None,
            "<<Price>>": None,
            "<<Sum_wo>>": None,
            "<<Sum_w>>": None,
        }

        if docxData[4]:
            ROW_TOKENS["<<Days>>"] = None

        TOTAL_TOKENS = {
            "<<TOTAL_WO>>": None,
            "<<TOTAL_W>>": None,
        }

        def _fmt_dec_comma(v: Decimal) -> str:
            x = v.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
            return f"{x:.2f}".replace(".", ",")


        def fmt_money_with_symbol(v) -> str:
            return symbCurrency + _fmt_dec_comma(Decimal(str(v)))


        def fmt_money_no_symbol(v) -> str:
            return _fmt_dec_comma(Decimal(str(v)))


        def _replace_in_paragraph_runs(paragraph, mapping: dict[str, str]) -> bool:
            changed = False
            for r in paragraph.runs:
                for k, v in mapping.items():
                    if k in r.text:
                        r.text = r.text.replace(k, v)
                        changed = True
            return changed


        def _replace_in_paragraph_fallback_merge(paragraph, mapping: dict[str, str]) -> bool:
            if not paragraph.runs:
                return False

            full = "".join(r.text for r in paragraph.runs)
            new = full
            for k, v in mapping.items():
                if k in new:
                    new = new.replace(k, v)

            if new == full:
                return False

            paragraph.runs[0].text = new
            for r in paragraph.runs[1:]:
                r.text = ""
            return True


        def _replace_in_paragraph(p, mapping: dict[str, str]) -> None:
            # Сначала пробуем точечную замену по runs, чтобы сохранить формат.
            _replace_in_paragraph_runs(p, mapping)
            # Если после этого остались токены (обычно они разбиты между runs),
            # выполняем fallback-слияние и замену.
            full_after_runs = "".join(r.text for r in p.runs)
            if any(token in full_after_runs for token in mapping):
                _replace_in_paragraph_fallback_merge(p, mapping)


        def _replace_in_table(table, mapping: dict[str, str]) -> None:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_in_paragraph(p, mapping)
                    for nested_table in cell.tables:
                        _replace_in_table(nested_table, mapping)


        def replace_placeholders_everywhere(doc: Document, mapping: dict[str, str]) -> None:
            def handle_paragraph(p):
                _replace_in_paragraph(p, mapping)

            for p in doc.paragraphs:
                handle_paragraph(p)

            for t in doc.tables:
                _replace_in_table(t, mapping)

            for sec in doc.sections:
                for hdrftr in (sec.header, sec.footer):
                    for p in hdrftr.paragraphs:
                        handle_paragraph(p)
                    for t in hdrftr.tables:
                        _replace_in_table(t, mapping)


        def _set_cell_text_keep_style(cell, text: str) -> None:
            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            if not p.runs:
                p.add_run(text)
            else:
                p.runs[0].text = text
                for r in p.runs[1:]:
                    r.text = ""
            for extra_p in cell.paragraphs[1:]:
                for r in extra_p.runs:
                    r.text = ""


        def _find_products_table(doc: Document):
            need = ["Наименование", "Каталожный товар", "Кол-во"]
            for t in doc.tables:
                if len(t.rows) == 0:
                    continue
                header = " | ".join(c.text for c in t.rows[0].cells)
                if all(x in header for x in need):
                    return t
            return doc.tables[1] if len(doc.tables) > 1 else None


        def _row_text_lower(row) -> str:
            return "|".join((c.text or "") for c in row.cells).lower()


        def _find_row_idx_by_token(table, token: str) -> int | None:
            tok = token.lower()
            for i, row in enumerate(table.rows):
                if tok in _row_text_lower(row):
                    return i
            return None

        def _find_header_row_idx(table) -> int:
            need = ["наименование", "каталожный", "кол-во"]
            for i, row in enumerate(table.rows):
                rt = _row_text_lower(row)
                if all(x in rt for x in need):
                    return i
            return 0

        def _find_total_row_idx(table) -> int:
            for i, row in enumerate(table.rows):
                for cell in row.cells:
                    if (cell.text or "").strip().lower() == "итого":
                        return i
            return len(table.rows) - 1

        def _fill_row_by_indices(row, idx: int, it: dict, sum_wo: Decimal, sum_w: Decimal):
            _set_cell_text_keep_style(row.cells[0], str(idx))
            _set_cell_text_keep_style(row.cells[1], it["name"])
            _set_cell_text_keep_style(row.cells[2], it["sku"])
            _set_cell_text_keep_style(row.cells[3], it["unit"])
            _set_cell_text_keep_style(row.cells[4], str(it["qty"]))
            _set_cell_text_keep_style(row.cells[5], Tools.formatPrice(_fmt_dec_comma(it["price"]), symbCurrency))
            _set_cell_text_keep_style(row.cells[6], Tools.formatPrice(_fmt_dec_comma(sum_wo), symbCurrency))
            _set_cell_text_keep_style(row.cells[7], Tools.formatPrice(_fmt_dec_comma(sum_w), symbCurrency))
            if docxData[4]:
                _set_cell_text_keep_style(row.cells[8], it["days"])


        def _fill_row_by_tokens(row, mapping: dict[str, str]):
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph_runs(p, mapping)
                    _replace_in_paragraph_fallback_merge(p, mapping)


        def fill_products_table(doc: Document, items: list[dict]) -> tuple[Decimal, Decimal, Decimal]:
            table = _find_products_table(doc)
            if table is None:
                raise RuntimeError("Не нашёл таблицу товаров в документе.")

            header_idx = _find_header_row_idx(table)
            total_idx = _find_total_row_idx(table)
            _normalize_products_table_headers(table, header_idx)

            template_idx = _find_row_idx_by_token(table, "<<Name>>")
            if template_idx is None:
                template_idx = _find_row_idx_by_token(table, "{{Name}}")

            if template_idx is None:
                template_idx = header_idx + 1
                if template_idx >= total_idx:
                    raise RuntimeError("В таблице нет строки-образца товара между заголовком и итогом.")

            template_tr = copy.deepcopy(table.rows[template_idx]._tr)

            for i in range(total_idx - 1, header_idx, -1):
                if i == template_idx:
                    continue
                table._tbl.remove(table.rows[i]._tr)

            total_idx = _find_total_row_idx(table)
            total_tr = table.rows[total_idx]._tr

            total_wo = Decimal("0.00")
            total_w = Decimal("0.00")

            template_row_text = _row_text_lower(table.rows[template_idx])
            use_tokens = ("<<Name>>" in template_row_text) or ("{{Name}}" in template_row_text)

            for n, it in enumerate(items, start=1):
                price = Decimal(str(it["price"]))
                qty = Decimal(str(it["qty"]))
                sum_wo = it['sum_wo']
                sum_w = it['sum_w']

                total_wo += sum_wo
                total_w += sum_w

                if n == 1:
                    row = table.rows[template_idx]
                else:
                    new_tr = copy.deepcopy(template_tr)
                    total_tr.addprevious(new_tr)
                    row = table.rows[_find_total_row_idx(table) - 1]

                if use_tokens:
                    row_map = {
                        "<<I>>": str(n),
                        "<<Name>>": it["name"],
                        "<<SKU>>": it["sku"],
                        "<<UNIT>>": it["unit"],
                        "<<QTY>>": str(it["qty"]),
                        "<<PRICE>>": Tools.num2text(fmt_money_with_symbol(price)),
                        "<<SUM_WO>>": currency[0] + _fmt_dec_comma(sum_wo),
                        "<<SUM_W>>": currency[0] + _fmt_dec_comma(sum_w),
                        "<<DAYS>>": it["days"],
                        "<<Days>>": it["days"],
                        "{{I}}": str(n),
                        "{{NAME}}": it["name"],
                        "{{SKU}}": it["sku"],
                        "{{UNIT}}": it["unit"],
                        "{{QTY}}": str(it["qty"]),
                        "{{PRICE}}": Tools.num2text(fmt_money_with_symbol(price)),
                        "{{SUM_WO}}": fmt_money_with_symbol(_fmt_dec_comma(sum_wo)),
                        "{{SUM_W}}": fmt_money_with_symbol(_fmt_dec_comma(sum_w)),
                    }
                    if docxData[4]:
                        row_map["{{DAYS}}"] = it["days"]
                        row_map["{{Days}}"] = it["days"]
                    _fill_row_by_tokens(row, row_map)
                else:
                    _fill_row_by_indices(row, n, it, sum_wo, sum_w)

            total_row = table.rows[_find_total_row_idx(table)]
            total_row_text = _row_text_lower(total_row)
            if any(token in total_row_text for token in ("<<total_wo>>", "<<total_w>>", "{{total_wo}}", "{{total_w}}")):
                total_wo_text = Tools.formatPrice(_fmt_dec_comma(total_wo), symbCurrency)
                total_w_text = Tools.formatPrice(_fmt_dec_comma(total_w), symbCurrency)
                totals_map = {
                    "<<TOTAL_WO>>": total_wo_text,
                    "<<TOTAL_W>>": total_w_text,
                    "{{TOTAL_WO}}": total_wo_text,
                    "{{TOTAL_W}}": total_w_text,
                    "<<total_wo>>": total_wo_text,
                    "<<total_w>>": total_w_text,
                    "{{total_wo}}": total_wo_text,
                    "{{total_w}}": total_w_text,
                }
                _fill_row_by_tokens(total_row, totals_map)
            else:
                _set_cell_text_keep_style(total_row.cells[6], currency[0] + _fmt_dec_comma(total_wo))
                _set_cell_text_keep_style(total_row.cells[7], currency[0] + _fmt_dec_comma(total_w))

            _set_row_bold(total_row)
            _optimize_products_table_layout(
                table,
                doc.sections[0],
                include_days=bool(docxData[4]),
                header_idx=header_idx,
            )

            vat_sum = (total_w - total_wo).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
            return total_wo, vat_sum, total_w


        def main():
            doc = Document(TEMPLATE_PATH)
            _apply_top_indent_for_multipage_table(
                doc,
                ITEMS,
                include_days=bool(docxData[4]),
            )

            total_wo, vat_sum, total_w = fill_products_table(doc, ITEMS)
            _insert_product_condition_after_delivery_terms(doc)
            mapping = dict(PLACEHOLDERS)
            mapping["<<TOTAL_WO_CNY>>"] = fmt_money_no_symbol(total_wo)
            mapping["<<VAT_CNY>>"] = fmt_money_no_symbol(round(vat_sum, 2))
            mapping["<<TOTAL_W_CNY>>"] = fmt_money_no_symbol(total_w)

            replace_placeholders_everywhere(doc, mapping)
            _normalize_offer_terms_paragraphs(doc)

            doc.save(OUTPUT_PATH)

        try:
            main()
            if output_format == "pdf":
                self.output_path = self._convert_docx_to_pdf(OUTPUT_PATH)
                try:
                    Path(OUTPUT_PATH).unlink(missing_ok=True)
                except Exception:
                    pass
            else:
                self.output_path = OUTPUT_PATH
            Tools.write_log("creating docx File...")
            Tools.write_log(f"saving docx to: {docx_output_dir}")
            self.success = True

        except Exception as e:
            Tools.write_log(f"Unnable to save Docx: {e}")
            self.error_message = str(e)

    @staticmethod
    def _normalize_output_format(value):
        normalized = str(value or "docx").strip().lower()
        return "pdf" if normalized == "pdf" else "docx"

    @staticmethod
    def _convert_docx_to_pdf(docx_path: str) -> str:
        source_path = Path(docx_path).expanduser().resolve()
        target_path = source_path.with_suffix(".pdf")

        for executable in ("soffice", "libreoffice"):
            command = shutil.which(executable)
            if not command:
                continue
            completed = subprocess.run(
                [
                    command,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(source_path.parent),
                    str(source_path),
                ],
                capture_output=True,
                text=True,
                timeout=120,
                check=False,
            )
            if completed.returncode == 0 and target_path.exists():
                return str(target_path)
            details = (completed.stderr or completed.stdout or "").strip()
            raise RuntimeError(f"Не удалось сохранить PDF через LibreOffice: {details}")

        try:
            from docx2pdf import convert
        except Exception as exc:
            raise RuntimeError(
                "Для сохранения КП в PDF нужен установленный LibreOffice "
                "или пакет docx2pdf с Microsoft Word."
            ) from exc

        convert(str(source_path), str(target_path))
        if not target_path.exists():
            raise RuntimeError("PDF не был создан после конвертации DOCX")
        return str(target_path)

class createExcelFile:
    FORMULA_EDITABLE_COLUMNS = (8, 9, 10, 11, 13)
    FORMULA_COLUMN_TO_EXCEL = {
        8: "I",
        9: "J",
        10: "K",
        11: "L",
        13: "N",
    }
    DEFAULT_FORMULAS = {
        8: "Custom*Logistic",
        9: "Customs/Amount",
        10: "UnitSalePrice*Markup",
        11: "RealPrice*Amount",
        13: "SupplierTerm+TermDelivery",
    }
    COLUMN_TITLES = {
        7: "Логистика",
        8: "Таможня",
        9: "Цена за ед.",
        10: "Цена реализации за ед. без НДС",
        11: "Итого реализации без НДС",
        13: "Срок поставки",
    }
    NAMED_VAR_PATTERN = re.compile(r"\$([^$]+)\$")
    TOKEN_PATTERN = re.compile(r"\b[A-Za-z_][A-Za-z0-9_]*\b")

    def __init__(self, data):
        self.output_path = ""
        self.success = False
        self.error_message = ""
        try:
            payload = self._normalize_payload(data)
            self._build_excel(payload)
            self.success = True
        except Exception as e:
            Tools.write_log(f"Unnable to save Excel: {e}")
            self.error_message = str(e)

    @staticmethod
    def _normalize_param_name(value):
        return str(value or "").strip().casefold()

    @staticmethod
    def _fmt_number(value):
        text = str(value).strip().replace(",", ".")
        if not text:
            return "0"
        try:
            num = float(text)
        except ValueError:
            return text
        if num.is_integer():
            return str(int(num))
        return f"{num:.10f}".rstrip("0").rstrip(".")

    @staticmethod
    def _currency_format(currency):
        if currency:
            return f'"{currency}"#,##0.00'
        return "#,##0.00"

    @staticmethod
    def _to_float(value, field_name):
        text = str(value or "").strip().replace(",", ".")
        if not text:
            raise ValueError(f'Поле "{field_name}" не заполнено')
        try:
            return float(text)
        except ValueError:
            return float(Tools.evalWithVars(text))

    @staticmethod
    def _wrap_round(formula, digits):
        expression = str(formula or "").strip()
        if expression.startswith("="):
            expression = expression[1:]
        return f"=ROUND({expression}, {digits})"

    @classmethod
    def _load_named_parameters(cls):
        params_data = Tools.load_json(Config.vars_path)
        parameters = {}
        for values in params_data.get("parameters", {}).values():
            if len(values) < 3:
                continue
            variable, value, calc_type = values[0], values[1], values[2]
            key = cls._normalize_param_name(variable)
            if not key:
                continue
            parameters[key] = (cls._fmt_number(value), str(calc_type))
        return parameters

    @classmethod
    def _normalize_named_parameters(cls, raw_parameters):
        normalized = {}
        if not isinstance(raw_parameters, dict):
            return normalized
        for key, value in raw_parameters.items():
            normalized_key = cls._normalize_param_name(key)
            if not normalized_key:
                continue
            if isinstance(value, (list, tuple)) and len(value) >= 2:
                param_value = value[0]
                calc_type = value[1]
            elif isinstance(value, dict):
                param_value = value.get("value", "")
                calc_type = value.get("calc_type", value.get("calc", ""))
            else:
                continue
            normalized[normalized_key] = (cls._fmt_number(param_value), str(calc_type))
        return normalized

    @staticmethod
    def _load_vat_multiplier():
        params_data = Tools.load_json(Config.vars_path)
        for values in params_data.get("parameters", {}).values():
            if len(values) < 3:
                continue
            name, value, calc_type = values[0], values[1], values[2]
            if str(name or "").strip().casefold() != "ндс":
                continue
            rate = float(str(value).replace(",", "."))
            if str(calc_type) == "percents":
                return 1 + rate / 100
            return 1 + rate
        return 1.0

    def _normalize_payload(self, data):
        if isinstance(data, dict):
            table_rows = [list(row) for row in data.get("table_rows", [])]
            request_number = str(
                data.get("request_number", Config.config.get("requestNumber", ""))
            ).strip()
            logistic_mode = int(data.get("logistic_mode", 0))
            logistic_value = self._to_float(data.get("logistic_value", 1), "Логистика")
            custom_value = self._to_float(data.get("custom_value", 1), "Таможня")
            markup_value = self._to_float(data.get("markup_value", 1), "Наценка")
            term_delivery = Tools.parse_int(data.get("term_delivery", 0), "Срок поставки", allow_zero=True)
            vat_multiplier = self._to_float(data.get("vat_multiplier", 1), "НДС")
            formula_expressions_raw = data.get("formula_expressions", {})
            formula_expressions = {}
            if isinstance(formula_expressions_raw, dict):
                for col, formulas in formula_expressions_raw.items():
                    try:
                        col_idx = int(col)
                    except (TypeError, ValueError):
                        continue
                    if col_idx not in self.FORMULA_EDITABLE_COLUMNS:
                        continue
                    if isinstance(formulas, list):
                        formula_expressions[col_idx] = [str(formula or "") for formula in formulas]
            logistic_formulas_raw = data.get("logistic_formulas", [])
            logistic_formulas = []
            if isinstance(logistic_formulas_raw, list):
                logistic_formulas = [str(formula or "").strip() for formula in logistic_formulas_raw]
            named_parameters = self._normalize_named_parameters(data.get("named_parameters", {}))
            if not named_parameters:
                named_parameters = self._load_named_parameters()
            docx_remote_url = str(data.get("docx_remote_url", "") or "").strip()
        else:
            table_rows = [list(row) for row in (data[0] if len(data) > 0 else [])]
            request_number = (
                str(data[4]).strip() if len(data) > 4 else str(Config.config.get("requestNumber", "")).strip()
            )
            legacy_params = data[1] if len(data) > 1 else (0, 1, 1)
            logistic_mode = int(legacy_params[0]) if len(legacy_params) > 0 else 0
            logistic_value = self._to_float(legacy_params[1] if len(legacy_params) > 1 else 1, "Логистика")
            custom_value = self._to_float(data[2] if len(data) > 2 else 1, "Таможня")
            markup_value = self._to_float(legacy_params[2] if len(legacy_params) > 2 else 1, "Наценка")
            term_delivery = Tools.parse_int(
                Config.config.get("termDelivery", "0"),
                "Срок поставки",
                allow_zero=True,
            )
            vat_multiplier = self._load_vat_multiplier()
            formula_expressions = {}
            logistic_formulas = []
            named_parameters = self._load_named_parameters()
            docx_remote_url = ""

        if not table_rows:
            raise ValueError("Нет данных для экспорта Excel")

        return {
            "table_rows": table_rows,
            "request_number": request_number,
            "logistic_mode": logistic_mode,
            "logistic_value": logistic_value,
            "custom_value": custom_value,
            "markup_value": markup_value,
            "term_delivery": term_delivery,
            "vat_multiplier": vat_multiplier,
            "logistic_formulas": logistic_formulas,
            "formula_expressions": formula_expressions,
            "named_parameters": named_parameters,
            "docx_remote_url": docx_remote_url,
        }

    def _row_value(self, row_values, preferred_index, fallback_index=None):
        if preferred_index is not None and preferred_index < len(row_values):
            return row_values[preferred_index]
        if fallback_index is not None and fallback_index < len(row_values):
            return row_values[fallback_index]
        return ""

    def _parse_row_number(self, row_values, row_index):
        raw_value = str(self._row_value(row_values, 0) or "").strip()
        if not raw_value:
            return row_index + 1
        try:
            return Tools.parse_int(raw_value, f"№ (строка {row_index + 1})", allow_zero=False)
        except ValueError:
            return row_index + 1

    def _parse_amount(self, row_values, row_index):
        raw_value = self._row_value(row_values, 4)
        return Tools.parse_int(raw_value, f"Кол-во (строка {row_index + 1})", allow_zero=False)

    def _parse_unit_price(self, row_values, row_index):
        raw_value = self._row_value(row_values, 5)
        currency, unit_price_text = Tools.parsePrice(raw_value)
        unit_price = Tools.parse_float(unit_price_text, f"Цена за ед. (строка {row_index + 1})", allow_zero=True)
        return currency, unit_price

    def _parse_supplier_term(self, row_values, row_index):
        raw_value = self._row_value(row_values, 14, fallback_index=7)
        try:
            return Tools.parse_delivery_days(raw_value)
        except ValueError:
            raise ValueError(f'Строка {row_index + 1}, столбец "Срок поставщика": некорректное значение')

    def _formula_for_row(self, formula_expressions, col, row_index):
        formulas = formula_expressions.get(col, [])
        if row_index < len(formulas):
            formula = str(formulas[row_index] or "").strip()
            if formula:
                return formula
        return self.DEFAULT_FORMULAS[col]

    def _formula_context_by_column(self, excel_row, payload):
        term_delivery_value = self._fmt_number(payload["term_delivery"])
        base = {
            "amount": f"E{excel_row}",
            "qty": f"E{excel_row}",
            "unitprice": f"F{excel_row}",
            "price": f"F{excel_row}",
            "totalprice": f"G{excel_row}",
            "logistic": f"H{excel_row}",
            "custom": self._fmt_number(payload["custom_value"]),
            "markup": self._fmt_number(payload["markup_value"]),
            "vat": self._fmt_number(payload["vat_multiplier"]),
            "supplierterm": f"O{excel_row}",
            "termdelivery": f"IF(O{excel_row}>0,{term_delivery_value},0)",
        }
        with_customs = dict(base)
        with_customs["customs"] = f"I{excel_row}"

        with_unit_sale = dict(with_customs)
        with_unit_sale["unitsaleprice"] = f"J{excel_row}"

        with_real_price = dict(with_unit_sale)
        with_real_price["realprice"] = f"K{excel_row}"

        with_totals = dict(with_real_price)
        with_totals["totalwithoutvat"] = f"L{excel_row}"
        with_totals["totalwithvat"] = f"M{excel_row}"

        return {
            8: dict(base),
            9: with_customs,
            10: with_unit_sale,
            11: with_real_price,
            13: with_totals,
        }

    def _replace_named_parameter(self, expression, named_parameters, row_index, col):
        def _replace(match):
            token = match.group(1).strip()
            key = self._normalize_param_name(token)
            if key not in named_parameters:
                raise ValueError(
                    f'Строка {row_index + 1}, столбец "{self.COLUMN_TITLES[col]}": '
                    f'неизвестная переменная "${token}$"'
                )
            value, calc_type = named_parameters[key]
            value_text = self._fmt_number(value)
            if calc_type == "percents":
                return f"({value_text})/100"
            if calc_type == "multiply":
                return f"*({value_text})"
            if calc_type == "division":
                return f"/({value_text})"
            return f"({value_text})"

        return self.NAMED_VAR_PATTERN.sub(_replace, expression)

    def _formula_to_excel(self, expression, context, named_parameters, row_index, col):
        formula_text = str(expression or "").strip().replace(",", ".")
        if formula_text.startswith("="):
            formula_text = formula_text[1:].strip()
        if not formula_text:
            raise ValueError(
                f'Строка {row_index + 1}, столбец "{self.COLUMN_TITLES[col]}": формула не может быть пустой'
            )

        formula_text = self._replace_named_parameter(formula_text, named_parameters, row_index, col).strip()
        while formula_text and formula_text[0] in "+*/":
            formula_text = formula_text[1:].strip()
        if not formula_text:
            raise ValueError(
                f'Строка {row_index + 1}, столбец "{self.COLUMN_TITLES[col]}": формула не может быть пустой'
            )

        def _replace_token(match):
            token = match.group(0)
            key = token.lower()
            if key not in context:
                key = key.replace("_", "")
            if key not in context:
                raise ValueError(
                    f'Строка {row_index + 1}, столбец "{self.COLUMN_TITLES[col]}": '
                    f'неизвестная переменная "{token}"'
                )
            return context[key]

        excel_expression = self.TOKEN_PATTERN.sub(_replace_token, formula_text)
        return f"={excel_expression}"

    REMOTE_LINK_LABEL_COLUMN = "D"
    REMOTE_LINK_URL_COLUMN = "E"

    @classmethod
    def _append_remote_link(
        cls,
        work_sheet,
        row: int,
        label: str,
        remote_url: str,
    ) -> None:
        url = str(remote_url or "").strip()
        if not url:
            return
        work_sheet[f"{cls.REMOTE_LINK_LABEL_COLUMN}{row}"] = label
        work_sheet[f"{cls.REMOTE_LINK_URL_COLUMN}{row}"] = url
        work_sheet[f"{cls.REMOTE_LINK_URL_COLUMN}{row}"].hyperlink = url
        work_sheet[f"{cls.REMOTE_LINK_URL_COLUMN}{row}"].style = "Hyperlink"

    @classmethod
    def _append_docx_remote_link(cls, work_sheet, row: int, remote_url: str) -> None:
        cls._append_remote_link(work_sheet, row, "Ссылка на КП DOCX", remote_url)

    @classmethod
    def _append_calculations_remote_link(cls, work_sheet, row: int, remote_url: str) -> None:
        cls._append_remote_link(work_sheet, row, "Ссылка на расчеты", remote_url)

    @classmethod
    def append_calculations_remote_link_to_file(
        cls,
        file_path: str,
        remote_url: str,
    ) -> None:
        workbook = load_workbook(file_path)
        try:
            worksheet = workbook.active
            target_row = None
            for row in range(1, worksheet.max_row + 1):
                labels = (
                    worksheet[f"{cls.REMOTE_LINK_LABEL_COLUMN}{row}"].value,
                    worksheet[f"A{row}"].value,
                )
                if any(str(label or "").strip() == "Ссылка на КП DOCX" for label in labels):
                    existing_url = ""
                    for column in (cls.REMOTE_LINK_URL_COLUMN, "B"):
                        value = str(worksheet[f"{column}{row}"].value or "").strip()
                        if value:
                            existing_url = value
                            break
                    cls._append_docx_remote_link(
                        worksheet,
                        row,
                        existing_url,
                    )
                    worksheet[f"A{row}"] = None
                    worksheet[f"B{row}"] = None
                    target_row = row + 1
                    break
            if target_row is None:
                target_row = worksheet.max_row + 1
            cls._append_calculations_remote_link(
                worksheet,
                target_row,
                remote_url,
            )
            workbook.save(file_path)
        finally:
            workbook.close()

    def _build_excel(self, payload):
        indent = int(Config.config["ExcelIndent"])
        request_number = payload["request_number"]
        today_date = datetime.now().strftime("%d_%m_%Y")
        documents_dir = Path.home() / "Documents"
        excel_output_dir = Tools.ensure_directory(
            Config.config.get("pathToSaveExcel") or Config.config.get("pathToSaveCP"),
            documents_dir,
        )
        Config.config["pathToSaveExcel"] = str(excel_output_dir)
        new_file_path = self.save_with_number(
            str(excel_output_dir / f"Расчеты_{request_number}_{today_date}_.xlsx")
        )
        self.output_path = new_file_path
        shutil.copy2(Config.template_path, new_file_path)
        wb = load_workbook(new_file_path)
        work_sheet = wb.active

        data_table = payload["table_rows"]
        first_data_row = 2 + indent
        last_data_row = len(data_table) + 1 + indent
        total_row = len(data_table) + 2 + indent
        logistic_value_text = self._fmt_number(payload["logistic_value"])
        vat_multiplier_text = self._fmt_number(payload["vat_multiplier"])
        currency = ""

        for row_index, row_values in enumerate(data_table):
            excel_row = row_index + first_data_row
            row_number = self._parse_row_number(row_values, row_index)
            amount = self._parse_amount(row_values, row_index)
            row_currency, unit_price = self._parse_unit_price(row_values, row_index)
            supplier_term = self._parse_supplier_term(row_values, row_index)
            if row_currency and not currency:
                currency = row_currency

            work_sheet[f"A{excel_row}"] = row_number
            work_sheet[f"B{excel_row}"] = self._row_value(row_values, 1)
            work_sheet[f"C{excel_row}"] = self._row_value(row_values, 2)
            work_sheet[f"D{excel_row}"] = self._row_value(row_values, 3)
            work_sheet[f"E{excel_row}"] = amount
            work_sheet[f"F{excel_row}"] = unit_price
            work_sheet[f"F{excel_row}"].number_format = self._currency_format(row_currency or currency)
            work_sheet[f"G{excel_row}"] = self._wrap_round(f"=F{excel_row}*E{excel_row}", 2)
            work_sheet[f"G{excel_row}"].number_format = self._currency_format(row_currency or currency)

            logistic_formula_source = ""
            if row_index < len(payload["logistic_formulas"]):
                logistic_formula_source = str(payload["logistic_formulas"][row_index] or "").strip()

            if logistic_formula_source:
                logistic_excel_formula = self._formula_to_excel(
                    logistic_formula_source,
                    self._formula_context_by_column(excel_row, payload)[8],
                    payload["named_parameters"],
                    row_index,
                    7,
                )
            else:
                if payload["logistic_mode"] == 1:
                    logistic_expression = (
                        f"IF(G{total_row}=0,0,G{excel_row}+{logistic_value_text}/G{total_row}*G{excel_row})"
                    )
                else:
                    logistic_expression = f"G{excel_row}*{logistic_value_text}"
                logistic_excel_formula = f"={logistic_expression}"

            work_sheet[f"H{excel_row}"] = self._wrap_round(logistic_excel_formula, 2)
            work_sheet[f"H{excel_row}"].number_format = self._currency_format(row_currency or currency)

            work_sheet[f"O{excel_row}"] = supplier_term
            work_sheet[f"O{excel_row}"].number_format = '0" дней"'

            contexts = self._formula_context_by_column(excel_row, payload)
            for col in self.FORMULA_EDITABLE_COLUMNS:
                formula_source = self._formula_for_row(payload["formula_expressions"], col, row_index)
                excel_formula = self._formula_to_excel(
                    formula_source,
                    contexts[col],
                    payload["named_parameters"],
                    row_index,
                    col,
                )
                if col == 13:
                    excel_formula = self._wrap_round(excel_formula, 0)
                else:
                    excel_formula = self._wrap_round(excel_formula, 2)

                cell_ref = f"{self.FORMULA_COLUMN_TO_EXCEL[col]}{excel_row}"
                work_sheet[cell_ref] = excel_formula
                if col in (8, 9, 10, 11):
                    work_sheet[cell_ref].number_format = self._currency_format(row_currency or currency)
                if col == 13:
                    work_sheet[cell_ref].number_format = '0" дней"'

            work_sheet[f"M{excel_row}"] = self._wrap_round(f"=L{excel_row}*{vat_multiplier_text}", 2)
            work_sheet[f"M{excel_row}"].number_format = self._currency_format(row_currency or currency)

        work_sheet[f"G{total_row}"] = f"=SUM(G{first_data_row}:G{last_data_row})"
        work_sheet[f"G{total_row}"].number_format = self._currency_format(currency)
        work_sheet[f"H{total_row}"] = f"=SUM(H{first_data_row}:H{last_data_row})"
        work_sheet[f"H{total_row}"].number_format = self._currency_format(currency)
        work_sheet[f"I{total_row}"] = f"=SUM(I{first_data_row}:I{last_data_row})"
        work_sheet[f"I{total_row}"].number_format = self._currency_format(currency)
        work_sheet[f"H{total_row + 1}"] = f"=H{total_row}-G{total_row}"
        work_sheet[f"H{total_row + 1}"].number_format = self._currency_format(currency)
        work_sheet[f"H{total_row + 2}"] = f"=H{total_row + 1}/G{total_row}"

        work_sheet[f"K{total_row}"] = "ИТОГО:"
        work_sheet[f"L{total_row}"] = f"=SUM(L{first_data_row}:L{last_data_row})"
        work_sheet[f"L{total_row}"].number_format = self._currency_format(currency)
        work_sheet[f"M{total_row}"] = f"=SUM(M{first_data_row}:M{last_data_row})"
        work_sheet[f"M{total_row}"].number_format = self._currency_format(currency)

        work_sheet[f"K{total_row + 3}"] = "Прибыль"
        work_sheet[f"K{total_row + 4}"] = "Прибыль %"
        work_sheet[f"L{total_row + 3}"] = f"=L{total_row}-I{total_row}"
        work_sheet[f"L{total_row + 3}"].number_format = self._currency_format(currency)
        work_sheet[f"L{total_row + 4}"] = f"=L{total_row + 3}/I{total_row}"
        work_sheet[f"L{total_row + 4}"].number_format = "0%"

        docx_remote_url = str(payload.get("docx_remote_url", "") or "").strip()
        if docx_remote_url:
            self._append_docx_remote_link(
                work_sheet,
                total_row + 5,
                docx_remote_url,
            )

        border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )

        for row in work_sheet.iter_rows():
            for cell in row:
                if self.cell_has_data(cell):
                    cell.border = border
                    cell.alignment = Alignment(horizontal="center", vertical="center")

        work_sheet.move_range("A1:O1", rows=indent)

        Tools.write_log("creating Excel File...")
        Tools.write_log(f"Excel path to save: {new_file_path}")
        Tools.write_log(f"Final path to save: {new_file_path}")
        wb.calculation.calcMode = "auto"
        wb.calculation.fullCalcOnLoad = True
        wb.calculation.forceFullCalc = True
        wb.save(new_file_path)
        try:
            if not force_excel_recalc(new_file_path):
                Tools.write_log(
                    "Excel formulas saved; automatic recalculation skipped."
                )
        except Exception as exc:
            Tools.write_log(f"Excel formulas saved but recalculation failed: {exc}")

    def cell_has_data(self, cell):
        if cell.value is not None and cell.value != "":
                return True
        return False

    def save_with_number(self, file_path):
        directory, filename = os.path.split(file_path)
        name, ext = os.path.splitext(filename)

        if not os.path.exists(file_path):
            return file_path

        counter = 1
        while True:
            new_filename = f"{name}({counter}){ext}"
            new_filepath = os.path.join(directory, new_filename)

            if not os.path.exists(new_filepath):
                return new_filepath

            counter += 1
